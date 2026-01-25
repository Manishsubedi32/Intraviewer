import io
from typing import Optional

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract text content from various file types.
    
    Supported formats:
    - PDF (.pdf)
    - Word Documents (.docx)
    - Plain Text (.txt)
    - Images (.png, .jpg, .jpeg) - using OCR
    """
    
    if not file_bytes:
        raise ValueError("Empty file provided")
    
    filename_lower = filename.lower()
    
    # ---- PDF FILES ----
    if filename_lower.endswith('.pdf'):
        return _extract_from_pdf(file_bytes)
    
    # ---- WORD DOCUMENTS ----
    elif filename_lower.endswith('.docx'):
        return _extract_from_docx(file_bytes)
    
    # ---- PLAIN TEXT ----
    elif filename_lower.endswith('.txt'):
        return _extract_from_txt(file_bytes)
    
    # ---- IMAGES (OCR) ----
    elif filename_lower.endswith(('.png', '.jpg', '.jpeg', '.webp', '.bmp', '.tiff')):
        return _extract_from_image(file_bytes)
    
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def _extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2 or pdfplumber."""
    text = ""
    
    # Try PyPDF2 first
    try:
        import PyPDF2
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        if text.strip():
            print(f"✅ PDF extracted with PyPDF2: {len(text)} chars")
            return text.strip()
    except ImportError:
        print("⚠️ PyPDF2 not installed")
    except Exception as e:
        print(f"⚠️ PyPDF2 failed: {e}")
    
    # Fallback to pdfplumber
    try:
        import pdfplumber
        pdf_file = io.BytesIO(file_bytes)
        
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if text.strip():
            print(f"✅ PDF extracted with pdfplumber: {len(text)} chars")
            return text.strip()
    except ImportError:
        print("⚠️ pdfplumber not installed")
    except Exception as e:
        print(f"⚠️ pdfplumber failed: {e}")
    
    # Fallback to pymupdf (fitz)
    try:
        import fitz  # PyMuPDF
        pdf_file = io.BytesIO(file_bytes)
        doc = fitz.open(stream=pdf_file, filetype="pdf")
        
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        
        doc.close()
        
        if text.strip():
            print(f"✅ PDF extracted with PyMuPDF: {len(text)} chars")
            return text.strip()
    except ImportError:
        print("⚠️ PyMuPDF not installed")
    except Exception as e:
        print(f"⚠️ PyMuPDF failed: {e}")
    
    # If all fail, try OCR on the PDF
    print("⚠️ All PDF text extractors failed, trying OCR...")
    return _extract_pdf_with_ocr(file_bytes)


def _extract_pdf_with_ocr(file_bytes: bytes) -> str:
    """Extract text from PDF using OCR (for scanned PDFs)."""
    try:
        import fitz  # PyMuPDF
        from PIL import Image
        import pytesseract
        
        text = ""
        pdf_file = io.BytesIO(file_bytes)
        doc = fitz.open(stream=pdf_file, filetype="pdf")
        
        for page_num, page in enumerate(doc):
            # Convert page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # OCR the image
            page_text = pytesseract.image_to_string(img)
            if page_text:
                text += page_text + "\n"
        
        doc.close()
        
        if text.strip():
            print(f"✅ PDF extracted with OCR: {len(text)} chars")
            return text.strip()
        
    except ImportError as e:
        print(f"⚠️ OCR dependencies missing: {e}")
    except Exception as e:
        print(f"⚠️ PDF OCR failed: {e}")
    
    raise ValueError("Failed to extract text from PDF. Install: pip install PyPDF2 pdfplumber PyMuPDF pytesseract pillow")


def _extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from Word documents."""
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
        
        if text.strip():
            print(f"✅ DOCX extracted: {len(text)} chars")
            return text.strip()
        else:
            raise ValueError("No text found in DOCX file")
            
    except ImportError:
        raise ValueError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def _extract_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text files."""
    try:
        # Try UTF-8 first
        text = file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        try:
            # Fallback to latin-1
            text = file_bytes.decode('latin-1')
        except Exception:
            # Last resort
            text = file_bytes.decode('utf-8', errors='ignore')
    
    if text.strip():
        print(f"✅ TXT extracted: {len(text)} chars")
        return text.strip()
    else:
        raise ValueError("Empty text file")


def _extract_from_image(file_bytes: bytes) -> str:
    """Extract text from images using OCR."""
    try:
        from PIL import Image
        import pytesseract
        
        # Open image
        img = Image.open(io.BytesIO(file_bytes))
        
        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Perform OCR
        text = pytesseract.image_to_string(img)
        
        if text.strip():
            print(f"✅ Image OCR extracted: {len(text)} chars")
            return text.strip()
        else:
            raise ValueError("No text found in image")
            
    except ImportError:
        raise ValueError("OCR dependencies missing. Run: pip install pytesseract pillow")
    except Exception as e:
        raise ValueError(f"Failed to extract text from image: {e}")